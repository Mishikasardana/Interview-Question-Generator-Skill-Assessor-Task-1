"""
DOCX text extraction module.

Purpose:
    Extract raw text from Word (.docx) resume files using ``python-docx``.
    Added alongside ``pdf_extractor.py`` so the resume upload flow accepts
    both PDF and DOCX — previously only PDF was accepted anywhere in the
    pipeline (file uploader UI, and this module).

Inputs:
    Path to a valid .docx file (``str`` or ``pathlib.Path``).

Outputs:
    Raw extracted text as a single string (may contain formatting artifacts).

Example usage:
    >>> from resume_processing.docx_extractor import extract_text_from_docx
    >>> text = extract_text_from_docx("resume.docx")
    >>> print(text[:200])

Assumptions:
    - Only modern .docx files are supported (the legacy binary .doc format
      is not — python-docx cannot read it).
    - Paragraph text and table cell text are both extracted; headers,
      footers, and text boxes are not (same MVP scope as the PDF extractor,
      which does not attempt OCR on scanned/image-only PDFs).
"""

from __future__ import annotations

from pathlib import Path
from typing import Union

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from resume_processing.exceptions import DocxExtractionError, FileValidationError

PathLike = Union[str, Path]

# .docx files are zip archives; this is the zip magic number (PK\x03\x04).
_DOCX_MAGIC_BYTES = b"PK\x03\x04"


def validate_docx_file(docx_path: PathLike) -> Path:
    """
    Validate that a path points to a readable, non-empty .docx file.

    Args:
        docx_path: Filesystem path to the candidate .docx file.

    Returns:
        Resolved absolute ``Path`` to the validated file.

    Raises:
        FileValidationError: If the path is missing, not a file, empty,
            or does not appear to be a .docx (zip-based) file.
    """
    path = Path(docx_path)

    if not path.exists():
        raise FileValidationError(f"Resume file not found: {path}")

    if not path.is_file():
        raise FileValidationError(f"Path is not a file: {path}")

    if path.stat().st_size == 0:
        raise FileValidationError(f"Resume file is empty: {path}")

    if path.suffix.lower() != ".docx":
        raise FileValidationError(
            f"Expected a .docx file, got '{path.suffix}' for: {path}"
        )

    try:
        header = path.read_bytes()[:4]
    except OSError as exc:
        raise FileValidationError(f"Cannot read resume file: {path}") from exc

    if not header.startswith(_DOCX_MAGIC_BYTES):
        raise FileValidationError(
            f"File does not appear to be a valid .docx: {path}. "
            "Note: the legacy .doc format is not supported, only .docx."
        )

    return path.resolve()


def extract_text_from_docx(docx_path: PathLike) -> str:
    """
    Extract all text from a .docx file (paragraphs and table cells).

    Args:
        docx_path: Filesystem path to the DOCX resume.

    Returns:
        Concatenated text from all paragraphs and table cells, joined with
        newline characters, in document order (paragraphs first, then
        tables — matching the simple, non-positional extraction already
        used for PDFs).

    Raises:
        FileValidationError: If the file does not exist or is not a valid
            .docx file.
        DocxExtractionError: If the file cannot be opened as a Word
            document, or contains no extractable text.
    """
    path = validate_docx_file(docx_path)

    try:
        document = Document(str(path))
    except PackageNotFoundError as exc:
        raise DocxExtractionError(
            f"Failed to open '{path.name}' as a Word document: {exc}"
        ) from exc
    except Exception as exc:
        raise DocxExtractionError(
            f"Unexpected error opening '{path.name}': {exc}"
        ) from exc

    chunks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            chunks.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    chunks.append(cell.text)

    if not chunks:
        raise DocxExtractionError(
            f"No extractable text found in '{path.name}'."
        )

    return "\n".join(chunks)
