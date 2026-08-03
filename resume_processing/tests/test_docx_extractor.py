"""Unit tests for docx_extractor.py."""

from __future__ import annotations

import pytest
from docx import Document

from resume_processing.exceptions import DocxExtractionError, FileValidationError
from resume_processing.docx_extractor import extract_text_from_docx, validate_docx_file


def _make_docx(tmp_path, filename: str = "resume.docx", with_table: bool = False):
    """Build a minimal real .docx file on disk for extraction tests."""
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Software Engineer with 5 years of experience.")
    if with_table:
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Python"
        table.rows[0].cells[1].text = "Docker"
    path = tmp_path / filename
    document.save(str(path))
    return path


def test_validate_docx_file_accepts_real_docx(tmp_path):
    path = _make_docx(tmp_path)
    validated = validate_docx_file(path)
    assert validated.exists()


def test_validate_docx_file_rejects_missing_file(tmp_path):
    with pytest.raises(FileValidationError):
        validate_docx_file(tmp_path / "does_not_exist.docx")


def test_validate_docx_file_rejects_empty_file(tmp_path):
    empty_path = tmp_path / "empty.docx"
    empty_path.write_bytes(b"")
    with pytest.raises(FileValidationError):
        validate_docx_file(empty_path)


def test_validate_docx_file_rejects_wrong_extension(tmp_path):
    path = _make_docx(tmp_path, filename="resume.pdf")
    with pytest.raises(FileValidationError):
        validate_docx_file(path)


def test_validate_docx_file_rejects_non_zip_content(tmp_path):
    fake_path = tmp_path / "fake.docx"
    fake_path.write_bytes(b"not a real docx file")
    with pytest.raises(FileValidationError):
        validate_docx_file(fake_path)


def test_extract_text_from_docx_returns_paragraph_text(tmp_path):
    path = _make_docx(tmp_path)
    text = extract_text_from_docx(path)
    assert "Jane Doe" in text
    assert "Software Engineer" in text


def test_extract_text_from_docx_includes_table_cells(tmp_path):
    path = _make_docx(tmp_path, with_table=True)
    text = extract_text_from_docx(path)
    assert "Python" in text
    assert "Docker" in text


def test_extract_text_from_docx_raises_on_no_extractable_text(tmp_path):
    document = Document()
    path = tmp_path / "empty_content.docx"
    document.save(str(path))
    with pytest.raises(DocxExtractionError):
        extract_text_from_docx(path)


def test_extract_text_from_docx_raises_on_corrupted_docx(tmp_path):
    # Valid zip magic bytes but not a real Word package.
    fake_path = tmp_path / "corrupt.docx"
    fake_path.write_bytes(b"PK\x03\x04" + b"not a real zip archive contents")
    with pytest.raises(DocxExtractionError):
        extract_text_from_docx(fake_path)
